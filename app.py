import io
import json

import streamlit as st
from streamlit_local_storage import LocalStorage

from utils.document_parser import parse_document
from utils.jd_fetcher import fetch_jd
from utils.cover_letter import generate_cover_letter
from utils.form_answers import generate_form_answers
from utils.docx_generator import generate_docx

localS = LocalStorage()

st.set_page_config(
    page_title="Job Application Agent",
    page_icon="\U0001F4C4",
    layout="wide",
)

st.title("\U0001F4C4 Job Application Agent")

# ---------------------------------------------------------------------------
# Session state defaults
# ---------------------------------------------------------------------------
for key, default in {
    "profile": {},
    "jd_data": None,
    "cover_letter": "",
    "form_answers": {},
    "original_cv_text": "",
    "tailored_cv_text": "",
    "tailored_cv_docx": None,
    "tailored_cv_pdf": None,
    "tailored_cv_changes": 0,
    "tailored_cv_rewrites": {},
    "cv_gap_analysis": None,
    "cv_tailor_file_bytes": None,
    "cv_tailor_file_name": "",
    "profile_loaded": False,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ---------------------------------------------------------------------------
# Load saved profile from localStorage
# ---------------------------------------------------------------------------
if not st.session_state["profile_loaded"]:
    saved_profile = localS.getItem("job_agent_profile")
    if saved_profile:
        try:
            st.session_state["profile"] = json.loads(saved_profile) if isinstance(saved_profile, str) else saved_profile
            st.session_state["profile_loaded"] = True
        except (json.JSONDecodeError, TypeError):
            pass

# ---------------------------------------------------------------------------
# Tabs
# ---------------------------------------------------------------------------
tab_profile, tab_application, tab_documents = st.tabs(
    ["My Profile", "New Application", "Documents"]
)

# ===== TAB 1: My Profile ====================================================
with tab_profile:
    st.header("Personal Information")

    _p = st.session_state.get("profile", {})

    col1, col2 = st.columns(2)
    with col1:
        full_name = st.text_input(
            "Full Name", value=_p.get("full_name", ""), key="pf_name"
        )
        email = st.text_input(
            "Email", value=_p.get("email", ""), key="pf_email"
        )
        phone = st.text_input(
            "Phone", value=_p.get("phone", ""), key="pf_phone"
        )
        linkedin = st.text_input(
            "LinkedIn",
            value=_p.get("linkedin", ""),
            placeholder="your-linkedin-username",
            key="pf_linkedin",
        )
        github = st.text_input(
            "GitHub",
            value=_p.get("github", ""),
            placeholder="your-github-username",
            key="pf_github",
        )
        address = st.text_input(
            "Address",
            value=_p.get("address", ""),
            key="pf_address",
            autocomplete="off",
        )
    with col2:
        city = st.text_input(
            "City", value=_p.get("city", ""), key="pf_city", autocomplete="off"
        )
        country = st.text_input(
            "Country",
            value=_p.get("country", ""),
            key="pf_country",
            autocomplete="off",
        )
        postal_code = st.text_input(
            "Postal Code",
            value=_p.get("postal_code", ""),
            key="pf_postal",
            autocomplete="off",
        )
        work_auth_options = ["Yes", "No", "Need sponsorship"]
        work_auth_idx = (
            work_auth_options.index(_p["work_auth"])
            if _p.get("work_auth") in work_auth_options
            else 0
        )
        work_auth = st.selectbox(
            "Work Authorization",
            work_auth_options,
            index=work_auth_idx,
            key="pf_work_auth",
        )

    st.divider()

    col3, col4 = st.columns(2)
    with col3:
        start_date = st.text_input(
            "Available Start Date",
            value=_p.get("start_date", ""),
            placeholder="e.g. May 1, 2026",
            key="pf_start",
        )
        notice_period = st.text_input(
            "Notice Period",
            value=_p.get("notice_period", ""),
            placeholder="e.g. Available immediately",
            key="pf_notice",
        )
    with col4:
        compensation = st.number_input(
            "Target Annual Compensation (\u20ac)",
            min_value=0,
            value=int(_p.get("compensation", 0)),
            step=1000,
            key="pf_comp",
        )
        referral_options = ["Company Website", "LinkedIn", "Referral", "Other"]
        referral_idx = (
            referral_options.index(_p["referral_source"])
            if _p.get("referral_source") in referral_options
            else 0
        )
        referral_source = st.selectbox(
            "How did you hear about this position?",
            referral_options,
            index=referral_idx,
            key="pf_referral",
        )
        if referral_source == "Other":
            referral_other = st.text_input(
                "Please specify",
                value=_p.get("referral_other", ""),
                placeholder="Please specify...",
                key="pf_referral_other",
            )
            if referral_other.strip():
                referral_source = referral_other.strip()

    col_save, col_clear = st.columns([1, 1])
    with col_save:
        if st.button("Save Profile", type="primary"):
            st.session_state["profile"] = {
                "full_name": full_name,
                "email": email,
                "phone": phone,
                "linkedin": linkedin,
                "github": github,
                "address": address,
                "city": city,
                "country": country,
                "postal_code": postal_code,
                "work_auth": work_auth,
                "start_date": start_date,
                "notice_period": notice_period,
                "compensation": compensation,
                "referral_source": referral_source,
            }
            localS.setItem(
                "job_agent_profile",
                json.dumps(st.session_state["profile"]),
            )
            st.success("Profile saved! Your information will be remembered next time.")
    with col_clear:
        if st.button("Clear Saved Profile"):
            localS.deleteItem("job_agent_profile")
            st.session_state["profile"] = {}
            st.session_state["profile_loaded"] = False
            st.rerun()

# ===== TAB 2: New Application ===============================================
with tab_application:

    # Helper to auto-save profile from widget values in Tab 1
    def _save_profile():
        st.session_state["profile"] = {
            "full_name": full_name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github,
            "address": address,
            "city": city,
            "country": country,
            "postal_code": postal_code,
            "work_auth": work_auth,
            "start_date": start_date,
            "notice_period": notice_period,
            "compensation": compensation,
            "referral_source": referral_source,
        }

    # ── STEP 1: Job Description ──────────────────────────────
    st.header("Step 1 \u2014 Job Description")

    jd_url = st.text_input("Job Description URL", key="jd_url")

    if st.button("Fetch JD"):
        if jd_url:
            with st.spinner("Fetching job description..."):
                jd_data = fetch_jd(jd_url)
            if jd_data.get("error"):
                st.warning(jd_data["error"])
                st.info("You can paste the job description manually below.")
            st.session_state["jd_data"] = jd_data
        else:
            st.warning("Please enter a URL first.")

    jd_manual = st.text_area(
        "Or paste the job description here (used if fetch fails or as supplement)",
        height=200,
        key="jd_manual",
    )

    if st.session_state["jd_data"] and not st.session_state["jd_data"].get("error"):
        jd = st.session_state["jd_data"]
        with st.expander("Fetched Job Description", expanded=True):
            if jd.get("title"):
                st.markdown(f"**Title:** {jd['title']}")
            if jd.get("company"):
                st.markdown(f"**Company:** {jd['company']}")
            if jd.get("location"):
                st.markdown(f"**Location:** {jd['location']}")
            if jd.get("requirements"):
                st.markdown("**Key Requirements:**")
                st.text(jd["requirements"][:1500])

    # Build effective JD dict (merge fetched + manual) for downstream steps
    _jd_dict = dict(st.session_state.get("jd_data") or {})
    if jd_manual.strip():
        existing = _jd_dict.get("full_text", "")
        _jd_dict["full_text"] = (existing + "\n\n" + jd_manual.strip()).strip()
    has_jd = bool(_jd_dict.get("full_text"))

    # ── STEP 2: Upload & Analyse CV ──────────────────────────
    if has_jd:
        st.divider()
        st.header("Step 2 \u2014 Upload & Analyse CV")

        tailor_cv_file = st.file_uploader(
            "Upload your CV (DOCX format required)",
            type=["docx"],
            key="tailor_cv_upload",
        )
        st.caption(
            "\U0001F4A1 Don't have a DOCX? In Google Docs: "
            "File \u2192 Download \u2192 Microsoft Word (.docx)"
        )

        if tailor_cv_file:
            tailor_cv_file.seek(0)
            st.session_state["cv_tailor_file_bytes"] = tailor_cv_file.read()
            st.session_state["cv_tailor_file_name"] = tailor_cv_file.name

        if st.button("Analyse CV vs Job Description", type="primary"):
            if not st.session_state.get("cv_tailor_file_bytes"):
                st.warning("Please upload your CV first.")
            else:
                file_like = io.BytesIO(st.session_state["cv_tailor_file_bytes"])
                file_like.name = st.session_state["cv_tailor_file_name"]
                cv_text = parse_document(file_like)

                with st.spinner("Analysing gaps between your CV and the job description..."):
                    from utils.cv_gap_analyzer import analyze_cv_gaps

                    gap_analysis = analyze_cv_gaps(cv_text, st.session_state["jd_data"])
                    st.session_state["cv_gap_analysis"] = gap_analysis
                    # Reset downstream state
                    st.session_state["tailored_cv_docx"] = None
                    st.session_state["tailored_cv_pdf"] = None
                    st.session_state["tailored_cv_text"] = None

        # Gap review & skill confirmation (shown after analysis)
        if st.session_state.get("cv_gap_analysis"):
            gap = st.session_state["cv_gap_analysis"]

            strengths = gap.get("existing_strengths", [])
            if strengths:
                with st.expander("Your existing strengths for this role", expanded=False):
                    for s in strengths:
                        st.markdown(f"- {s}")

            suggestions = gap.get("rewrite_suggestions", [])
            if suggestions:
                with st.expander("How we'll improve your bullet points", expanded=False):
                    for s in suggestions:
                        st.markdown(f"- {s}")

            missing_tools = gap.get("missing_tools", [])
            missing_skills = gap.get("missing_skills", [])
            all_missing = missing_tools + missing_skills

            confirmed_additions = []

            if all_missing:
                st.markdown(
                    "**The job description requires the following that aren't in your CV.**"
                )
                st.markdown(
                    "Check the ones you actually have experience with \u2014 "
                    "we'll add them to your CV:"
                )
                st.markdown("")

                for item in all_missing:
                    name = item.get("name", "")
                    context = item.get("jd_context", "")
                    category = item.get("category", "skill")
                    icon = "\U0001F527" if category == "tool" else "\U0001F4A1"

                    checked = st.checkbox(
                        f"{icon} **{name}** \u2014 *{context}*",
                        key=f"skill_confirm_{name.replace(' ', '_')}",
                        value=False,
                    )
                    if checked:
                        confirmed_additions.append(name)
            else:
                st.success("Your CV already covers the key requirements for this role.")

            # ── STEP 3: Generate Tailored CV ─────────────────────
            st.divider()
            st.header("Step 3 \u2014 Generate Tailored CV")

            if confirmed_additions:
                st.info(f"Will add to your CV: {', '.join(confirmed_additions)}")

            if st.button("Generate Tailored CV", type="primary"):
                _save_profile()

                cv_bytes = st.session_state["cv_tailor_file_bytes"]

                from utils.cv_xml_tailor import tailor_cv_xml_with_additions

                with st.spinner(
                    "Rewriting bullet points and applying your confirmed skills..."
                ):
                    docx_out, pdf_out, changes, rewrites = (
                        tailor_cv_xml_with_additions(
                            cv_bytes,
                            st.session_state["jd_data"],
                            st.session_state["profile"],
                            confirmed_additions,
                        )
                    )
                st.session_state["tailored_cv_docx"] = docx_out
                st.session_state["tailored_cv_pdf"] = pdf_out
                st.session_state["tailored_cv_changes"] = changes
                st.session_state["tailored_cv_rewrites"] = rewrites
                st.session_state["tailored_cv_text"] = None
                st.success(f"Done \u2014 {changes} sections updated.")

            # Tailored CV output
            if st.session_state.get("tailored_cv_docx"):
                applicant_name = (
                    st.session_state.get("profile", {})
                    .get("full_name", "CV")
                    .replace(" ", "_")
                )
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    st.download_button(
                        "Download Tailored CV (DOCX)",
                        data=st.session_state["tailored_cv_docx"],
                        file_name=f"Tailored_CV_{applicant_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                with col_dl2:
                    if st.session_state.get("tailored_cv_pdf"):
                        st.download_button(
                            "Download Tailored CV (PDF)",
                            data=st.session_state["tailored_cv_pdf"],
                            file_name=f"Tailored_CV_{applicant_name}.pdf",
                            mime="application/pdf",
                        )
                    else:
                        st.caption(
                            "PDF: install LibreOffice or convert DOCX manually."
                        )

                if st.session_state.get("tailored_cv_changes", 0) > 0:
                    with st.expander("See what changed"):
                        for idx, new_text in st.session_state.get(
                            "tailored_cv_rewrites", {}
                        ).items():
                            st.markdown(f"**[{idx}]** {new_text}")

    # ── STEP 4: Generate Cover Letter ────────────────────────
    if has_jd:
        st.divider()
        st.header("Step 4 \u2014 Generate Cover Letter")

        app_type = st.selectbox(
            "Application Type",
            ["Internship", "Working Student", "Praedoc", "Full-time"],
            key="app_type",
        )

        if st.button("Generate Cover Letter", type="primary"):
            _save_profile()
            profile = st.session_state["profile"]
            if not profile.get("full_name"):
                st.warning("Please save your profile first (My Profile tab).")
            elif not st.session_state.get("cv_tailor_file_bytes"):
                st.warning("Please upload your CV in Step 2 first.")
            else:
                with st.spinner("Preparing documents..."):
                    # Use tailored CV if available, otherwise original upload
                    if st.session_state.get("tailored_cv_docx"):
                        from docx import Document

                        doc = Document(io.BytesIO(st.session_state["tailored_cv_docx"]))
                        cv_text = "\n".join(
                            p.text for p in doc.paragraphs if p.text.strip()
                        )
                    else:
                        file_like = io.BytesIO(st.session_state["cv_tailor_file_bytes"])
                        file_like.name = st.session_state["cv_tailor_file_name"]
                        cv_text = parse_document(file_like)

                    texts = ["=== CV ===\n" + cv_text]
                    # Read additional docs from session state (uploaders rendered below)
                    for tr in st.session_state.get("tr_upload") or []:
                        tr.seek(0)
                        texts.append("=== Transcript ===\n" + parse_document(tr))
                    enroll = st.session_state.get("enroll_upload")
                    if enroll:
                        enroll.seek(0)
                        texts.append(
                            "=== Enrollment Certificate ===\n"
                            + parse_document(enroll)
                        )
                    for od in st.session_state.get("other_upload") or []:
                        od.seek(0)
                        texts.append(
                            "=== Other Document ===\n" + parse_document(od)
                        )
                    applicant_text = "\n\n".join(texts)

                with st.spinner("Generating cover letter..."):
                    cover_letter = generate_cover_letter(
                        applicant_text, _jd_dict, app_type, profile
                    )
                    st.session_state["cover_letter"] = cover_letter

                st.success("Cover letter generated!")

        # Cover letter output
        if st.session_state["cover_letter"]:
            edited_letter = st.text_area(
                "Edit your cover letter below:",
                value=st.session_state["cover_letter"],
                height=400,
                key="cl_edit",
            )

            profile = st.session_state.get("profile", {})
            applicant_name = profile.get("full_name", "cover_letter").replace(" ", "_")

            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                docx_buf = generate_docx(edited_letter, profile, _jd_dict)
                st.download_button(
                    "Download Cover Letter as DOCX",
                    data=docx_buf,
                    file_name=f"Cover_Letter_{applicant_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with col_dl2:
                from utils.pdf_generator import generate_pdf

                pdf_buf = generate_pdf(edited_letter, profile, _jd_dict)
                st.download_button(
                    "Download Cover Letter as PDF",
                    data=pdf_buf,
                    file_name=f"Cover_Letter_{applicant_name}.pdf",
                    mime="application/pdf",
                )

            if st.session_state.get("jd_data"):
                jd = st.session_state["jd_data"]
                title = jd.get("title", "Application")
                contact_email = jd.get("contact_email", "")

                if contact_email:
                    import urllib.parse
                    subject = urllib.parse.quote(f"Application \u2013 {title}")
                    body = urllib.parse.quote(
                        "Dear Hiring Team,\n\nPlease find my application "
                        "documents attached.\n\nBest regards,\n"
                        + profile.get("full_name", "")
                    )
                    mailto_href = f"mailto:{contact_email}?subject={subject}&body={body}"
                    st.markdown(f'**Contact:** [{contact_email}]({mailto_href})')
                else:
                    st.info(
                        "No contact email found in job description. "
                        "Check the JD page manually."
                    )

        st.divider()
        st.subheader("Additional Documents (optional)")
        col_up1, col_up2 = st.columns(2)
        with col_up1:
            transcripts = st.file_uploader(
                "Transcripts", type=["pdf"], accept_multiple_files=True, key="tr_upload"
            )
        with col_up2:
            enrollment = st.file_uploader(
                "Enrollment Certificate", type=["pdf"], key="enroll_upload"
            )
            other_docs = st.file_uploader(
                "Other Documents",
                type=["pdf"],
                accept_multiple_files=True,
                key="other_upload",
            )

    # ── STEP 5: Form Answers ─────────────────────────────────
    if st.session_state["cover_letter"]:
        st.divider()
        st.header("Step 5 \u2014 Form Answers")

        extra_questions = st.text_area(
            "Paste additional job-specific form questions (optional)",
            height=100,
            key="extra_qs",
        )

        if st.button("Generate Form Answers", type="primary"):
            _save_profile()
            profile = st.session_state["profile"]

            # Build applicant text same as Step 4
            if st.session_state.get("tailored_cv_docx"):
                from docx import Document

                doc = Document(io.BytesIO(st.session_state["tailored_cv_docx"]))
                cv_text = "\n".join(
                    p.text for p in doc.paragraphs if p.text.strip()
                )
            else:
                file_like = io.BytesIO(st.session_state["cv_tailor_file_bytes"])
                file_like.name = st.session_state["cv_tailor_file_name"]
                cv_text = parse_document(file_like)

            applicant_text = "=== CV ===\n" + cv_text

            with st.spinner("Generating form answers..."):
                answers = generate_form_answers(
                    applicant_text, _jd_dict, profile, extra_questions
                )
                st.session_state["form_answers"] = answers

            st.success("Form answers generated!")

        if st.session_state["form_answers"]:
            for key, value in st.session_state["form_answers"].items():
                st.markdown(f"**{key}**")
                st.code(value, language=None)

# ===== TAB 3: Documents =====================================================
with tab_documents:
    st.header("Uploaded Documents")

    all_files = []
    cv_bytes = st.session_state.get("cv_tailor_file_bytes")
    cv_name = st.session_state.get("cv_tailor_file_name", "")
    if cv_bytes:
        all_files.append(("CV (DOCX)", cv_name, cv_bytes))
    if st.session_state.get("tailored_cv_docx"):
        applicant_name = (
            st.session_state.get("profile", {})
            .get("full_name", "CV")
            .replace(" ", "_")
        )
        all_files.append((
            "Tailored CV",
            f"Tailored_CV_{applicant_name}.docx",
            st.session_state["tailored_cv_docx"],
        ))
    for i, tr in enumerate(transcripts if "transcripts" in dir() else []):
        tr.seek(0)
        all_files.append((f"Transcript {i + 1}", tr.name, tr.read()))
    if "enrollment" in dir() and enrollment:
        enrollment.seek(0)
        all_files.append(("Enrollment Certificate", enrollment.name, enrollment.read()))
    for i, od in enumerate(other_docs if "other_docs" in dir() else []):
        od.seek(0)
        all_files.append((f"Other Document {i + 1}", od.name, od.read()))

    if not all_files:
        st.info("No documents uploaded yet. Go to 'New Application' to upload.")
    else:
        for label, fname, fbytes in all_files:
            col_a, col_b, col_c = st.columns([3, 1, 1])
            with col_a:
                st.markdown(f"**{label}:** {fname}")
            with col_b:
                size_kb = len(fbytes) / 1024
                if size_kb >= 1024:
                    st.caption(f"{size_kb / 1024:.1f} MB")
                else:
                    st.caption(f"{size_kb:.0f} KB")
            with col_c:
                st.download_button(
                    "Download",
                    data=fbytes,
                    file_name=fname,
                    key=f"dl_{label}_{fname}",
                )

    st.divider()
    st.caption(
        "Note: Upload documents directly to the application portal separately."
    )
