import io
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _extract_username(value: str, domain: str) -> tuple:
    """Returns (display_text, full_url) regardless of input format."""
    value = value.strip().rstrip("/")
    if value.startswith("http"):
        parts = value.split("/")
        username = parts[-1] if parts[-1] else parts[-2]
        return username, value
    else:
        if domain == "linkedin":
            return value, f"https://www.linkedin.com/in/{value}"
        else:
            return value, f"https://github.com/{value}"


def _add_hyperlink_line(doc, parts):
    """
    parts: list of (display_text, url) or (plain_text, None)
    Adds a paragraph with mixed hyperlinks and plain text separated by " | "
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, (text, url) in enumerate(parts):
        if i > 0:
            run = para.add_run(" | ")
            run.font.name = "Arial"
            run.font.size = Pt(11)

        if url:
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(
                qn("r:id"),
                doc.part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                ),
            )
            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "Hyperlink")
            rPr.append(rStyle)
            new_run.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            new_run.append(t)
            hyperlink.append(new_run)
            para._p.append(hyperlink)
        else:
            run = para.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)

    return para


def _add_line(doc, text, font_size=Pt(11), bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = font_size
    run.bold = bold
    return para


def _add_blank(doc):
    doc.add_paragraph()


def generate_docx(cover_letter_text: str, profile_dict: dict, jd_dict: dict) -> io.BytesIO:
    doc = Document()

    # A4 page, 0.8-inch margins for single-page fit
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    name = profile_dict.get("full_name", "")
    email = profile_dict.get("email", "")
    phone = profile_dict.get("phone", "")
    city = profile_dict.get("city", "")
    country = profile_dict.get("country", "")
    postal_code = profile_dict.get("postal_code", "")
    linkedin = profile_dict.get("linkedin", "")
    github = profile_dict.get("github", "")

    # --- 1. HEADER BLOCK ---
    _add_line(doc, name, font_size=Pt(13), bold=True)

    # Line 2: postal, city, country
    location_parts = []
    if postal_code:
        location_parts.append(postal_code)
    if city:
        location_parts.append(city)
    if country:
        location_parts.append(country)
    if location_parts:
        _add_line(doc, ", ".join(location_parts))

    # Line 3: phone | email (email as mailto hyperlink)
    contact_parts = []
    if phone:
        contact_parts.append((phone, None))
    if email:
        contact_parts.append((email, f"mailto:{email}"))
    if contact_parts:
        _add_hyperlink_line(doc, contact_parts)

    # Line 4: linkedin | github (clickable)
    link_parts = []
    if linkedin:
        display, url = _extract_username(linkedin, "linkedin")
        link_parts.append((f"linkedin.com/in/{display}", url))
    if github:
        display, url = _extract_username(github, "github")
        link_parts.append((f"github.com/{display}", url))
    if link_parts:
        _add_hyperlink_line(doc, link_parts)

    _add_blank(doc)

    # --- 2. DATE ---
    _add_line(doc, date.today().strftime("%B %d, %Y"))
    _add_blank(doc)

    # --- 3. RECIPIENT BLOCK ---
    company = jd_dict.get("company", "")
    department = jd_dict.get("department", "")
    location = jd_dict.get("location", "")

    if company:
        _add_line(doc, company, bold=True)
    if department:
        _add_line(doc, department)
    if location:
        _add_line(doc, location)

    _add_blank(doc)

    # --- 4. SUBJECT LINE ---
    title = jd_dict.get("title", "")
    if title:
        _add_line(doc, f"Re: {title}", bold=True)
        _add_blank(doc)

    # --- 5. SALUTATION ---
    _add_line(doc, "Dear Hiring Team,")
    _add_blank(doc)

    # --- 6. BODY ---
    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        para_text = para_text.replace("\n", " ")
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(para_text)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    # --- 7. CLOSING ---
    _add_blank(doc)
    _add_line(doc, "Best regards,")
    _add_blank(doc)
    _add_line(doc, name, bold=True)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
