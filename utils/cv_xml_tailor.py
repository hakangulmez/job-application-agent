import io
import json
import os
import re
import subprocess
import tempfile

import anthropic
import streamlit as st
from docx import Document


def extract_paragraphs(docx_bytes: bytes) -> list:
    """Extract paragraphs with their text and position index."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name,
            })
    return paragraphs


def _apply_rewrites(docx_bytes: bytes, rewrites: dict) -> tuple:
    """
    Apply paragraph rewrites to a DOCX, preserving mixed bold/normal formatting.
    Returns (Document, changes_made).
    """
    doc = Document(io.BytesIO(docx_bytes))
    changes_made = 0

    for i, para in enumerate(doc.paragraphs):
        if i not in rewrites:
            continue
        if not para.runs:
            para.text = rewrites[i]
            changes_made += 1
            continue

        # Check if paragraph has mixed bold formatting
        # (some runs bold, some not — e.g. "Data & Tools:" bold + rest normal)
        bold_values = set(
            bool(run.bold) for run in para.runs if run.text.strip()
        )
        has_mixed_bold = len(bold_values) > 1

        if has_mixed_bold:
            # Collect bold prefix and normal text
            bold_prefix = ""
            for run in para.runs:
                if run.bold:
                    bold_prefix += run.text
                else:
                    break  # stop at first non-bold run

            # Clear all runs
            for run in para.runs:
                run.text = ""

            # Restore: first run gets bold prefix
            para.runs[0].text = bold_prefix
            para.runs[0].bold = True

            # Strip the bold prefix from the rewritten text to get the new tail
            new_text = rewrites[i]
            if new_text.startswith(bold_prefix):
                new_tail = new_text[len(bold_prefix):]
            else:
                # Claude may have slightly altered the prefix; use full text
                new_tail = new_text.replace(bold_prefix.strip(), "", 1).strip()
                if new_tail == new_text:
                    # Prefix not found at all — just put everything in first run
                    para.runs[0].text = new_text
                    changes_made += 1
                    continue

            # Find first non-bold run for the tail text
            placed = False
            for run in para.runs[1:]:
                if not run.bold or run.bold is None:
                    run.text = new_tail
                    run.bold = False
                    placed = True
                    break

            if not placed and len(para.runs) > 1:
                # All remaining runs are bold — use the second one as normal
                para.runs[1].text = new_tail
                para.runs[1].bold = False
            elif not placed:
                # Only one run exists — put full rewritten text there
                para.runs[0].text = new_text
        else:
            # Simple case: uniform formatting
            first_run = para.runs[0]
            bold = first_run.bold
            italic = first_run.italic
            font_size = first_run.font.size
            font_name = first_run.font.name
            for run in para.runs:
                run.text = ""
            first_run.text = rewrites[i]
            first_run.bold = bold
            first_run.italic = italic
            first_run.font.size = font_size
            first_run.font.name = font_name

        changes_made += 1

    return doc, changes_made


def tailor_cv_xml_with_additions(
    docx_bytes: bytes,
    jd_dict: dict,
    profile_dict: dict,
    confirmed_additions: list,
) -> tuple:
    """
    Rewrite CV bullet points in-place, preserving original DOCX formatting, and incorporate confirmed skill additions.
    confirmed_additions: list of skill/tool names user confirmed they have.
    """
    paragraphs = extract_paragraphs(docx_bytes)

    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    jd_text = jd_dict.get("full_text", "")[:2000]

    para_list = "\n".join(
        f"[{p['index']}] {p['text']}" for p in paragraphs
    )

    confirmed_str = ""
    if confirmed_additions:
        items = "\n".join(f"- {s}" for s in confirmed_additions)
        confirmed_str = f"""
CONFIRMED ADDITIONS (user has verified they have experience with these):
{items}

You MAY naturally incorporate these into relevant bullet points where they fit.
For example if "Databricks" is confirmed and a bullet mentions data pipelines,
you can add "using Databricks" to that bullet.
Only add them where they fit naturally — do not force them into unrelated bullets.
"""

    system = f"""You are a CV tailoring assistant. Rewrite CV bullet points to match a job description.

Rules:
- ABSOLUTE RULE: Do NOT add any tool, technology, or skill NOT in the original
  paragraph UNLESS it appears in the CONFIRMED ADDITIONS list below.
- Reframe existing content using JD keywords naturally
- Keep all facts accurate — never invent metrics or experiences
- Do not rewrite section headers, institution names, dates, or contact info
- Focus rewrites on bullet points and descriptions
- For Skills lines that start with a bold label like 'Data & Tools:' or \
'Programming:', keep the label exactly as-is and only modify the list \
of items after the colon.
- CRITICAL: Do NOT inject any domain, industry, or sector keywords \
(e.g. 'climate', 'fintech', 'healthcare') from the JD into the CV \
unless they appear in the CONFIRMED ADDITIONS list. \
Only reframe using the candidate's own existing language.

{confirmed_str}

Respond with JSON only:
{{"rewrites": {{"INDEX": "rewritten text"}}}}
Where INDEX is the paragraph index number.
Only include paragraphs that genuinely benefit from rewriting."""

    user_msg = f"""JOB DESCRIPTION:
{jd_text}

CV PARAGRAPHS:
{para_list}

Return JSON with rewrites."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=2000,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )

    response_text = response.content[0].text
    json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
    rewrites = {}
    if json_match:
        try:
            data = json.loads(json_match.group())
            rewrites = {int(k): v for k, v in data.get("rewrites", {}).items()}
        except (json.JSONDecodeError, ValueError):
            rewrites = {}

    doc, changes_made = _apply_rewrites(docx_bytes, rewrites)

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes_out = docx_buf.getvalue()
    pdf_bytes = docx_to_pdf(docx_bytes_out)

    return docx_bytes_out, pdf_bytes, changes_made, rewrites


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        soffice_candidates = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "soffice",
            "libreoffice",
        ]

        soffice = None
        for candidate in soffice_candidates:
            try:
                result = subprocess.run(
                    [candidate, "--version"],
                    capture_output=True,
                    timeout=5,
                )
                if result.returncode == 0:
                    soffice = candidate
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

        if soffice:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True,
                timeout=60,
            )

            pdf_path = os.path.join(tmpdir, "input.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read()

        return b""
